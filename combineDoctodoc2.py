import openai
import dotenv
import os
from docx import Document
from docx.shared import Inches
import subprocess
import pandas as pd
import tempfile
import shutil
import re
from extraction import extract_tables_and_images

dotenv.load_dotenv()
openai.api_key = os.getenv("OPENAI_API_KEY")


# Function to convert DOCX to Markdown and extract media
def convert_docx_to_md_and_extract_media(docx_file_path, media_path):
    pandoc_path = os.getenv('PANDOC_PATH')
    if pandoc_path is None:
        raise ValueError("PANDOC_PATH environment variable is not set.")

    if not os.path.exists(media_path):
        os.makedirs(media_path)

    output_md_path = os.path.join(media_path, "temp.md")
    try:
        subprocess.run([pandoc_path, "-s", docx_file_path, "-t", "markdown",
                        "--extract-media", media_path, "-o", output_md_path], check=True)
    except subprocess.CalledProcessError as e:
        raise ValueError(f"Error during conversion: {e}")
    return output_md_path


def redefine_from_markdown(markdown_content):
    style_prompt = """
Your task is to ensure uniformity in the provided document while preserving its original style and tone. Begin by closely analyzing the initial sentences and paragraphs to understand the nuances of the writing style, including the usage of adverbs, adjectives, sentence structure, and word choice. Maintain the format and structure of the initial text, including the number of paragraphs. Incorporate these style elements throughout the entire document without altering the format.
Adjust language and wording as necessary to maintain coherence, but ensure that the structure and format of the text remain consistent with the original. Pay careful attention to the voice and perspective established in the beginning and strive to emulate it consistently. Your objective is to create a seamless transition between the original content and the converted format while preserving the essence and personality conveyed in the text.
Please provide the output in plain text format without any Markdown.

"""
# 169 tokens style_prompt
    response = openai.chat.completions.create(model="gpt-4",
        messages=[
            {"role": "system", "content": style_prompt},
            {"role": "user", "content": markdown_content}
        ],
        max_tokens=4000,
        temperature=0.3)

    extracted_elements = response.choices[0].message.content

    if isinstance(extracted_elements, list):
        extracted_elements = " ".join(extracted_elements)

    return extracted_elements


def process_text_with_prompt(text):
    """
    Process the extracted text with a structured prompt for formatting and style adjustments.
    """
    instructional_part = """Instructions
Given the text provided, your task is to transform it into a well-structured document format. Carefully assess the initial sentences to grasp the tone, writing style, and overarching format. Your primary objective is to extend these attributes throughout the entire document, ensuring that it maintains a consistent tone, style, and format as though it was crafted by a single author.
Focus on the following aspects to create a coherent and unified document:
- Core Elements Identification: Determine the main theme to establish an apt title and use any mentioned topics or keywords as headings or subheadings to logically organize the content.
- Uniform Tone and Style: Extend the initial tone and style across the document, ensuring smooth transitions between sections and preserving the text's essence and personality.
- Coherence and Consistency: Pay close attention to sentence structure, word choice, and overall coherence to ensure the document feels naturally consistent.
- Document Structure: Ensure the text is well-organized and properly aligned, maintaining the original flow, tone, and writing style, adjusting only for clarity and coherence.
- Tables: For any placeholder in the format "Table: Table_{i}", the placeholder should be retained as-is and included at the exact position it occupied in the original text.
- Images: Similarly, for image placeholders in the format "Image: image{i}", these should also be kept unchanged and retained at their original positions in the output.
Final Output: Your reformatted document should be presented in Markdown format. This directive aims to enhance readability and clarity, adhering to the document's structured requirements. The essence of the original text should be evident, with modifications serving to improve the organization and presentation.
Do not include these instructions in your output.
Instructions End."""

    content_part = f"Content to Process:\n{text}\nContent End."

    full_prompt = f"{instructional_part}\n\n{content_part}"

    response = openai.chat.completions.create(
        model="gpt-4",
        messages=[
            {"role": "system", "content": "You are a document conversion system. Your task is to convert the provided text into a uniform formatted document while preserving the tone and writing style consistent with the original content."},
            {"role": "user", "content": full_prompt}
        ],
        max_tokens=4000,
        temperature=0.7
    )

    second_extracted_elements = response.choices[0].message.content
    # print(extracted_elements)

    if isinstance(second_extracted_elements, list):
        second_extracted_elements = " ".join(second_extracted_elements)

    return second_extracted_elements

# def markdown_to_docx(markdown_content, output_docx_path, images):
#     doc = Document()
#     lines = markdown_content.split('\n')
#     image_idx = 0

#     for line in lines:
#         if line.startswith("Image:"):
#             img_caption = line.split(":")[1].strip()
#             if image_idx < len(images):
#                 img_path = images[image_idx]
#                 # Add image to the document
#                 doc.add_picture(img_path, width=Inches(5))
#                 # Add caption below the image
#                 last_paragraph = doc.paragraphs[-1]
#                 last_paragraph.alignment = 1  # Center alignment
#                 doc.add_paragraph(img_caption, style='Caption')
#                 image_idx += 1
#         elif line.startswith("Table:"):
#             doc.add_paragraph(line)
#         else:
#             doc.add_paragraph(line)


#     # Save the document with images and table placeholders
#     temp_docx_path = "temp_with_images.docx"
#     doc.save(temp_docx_path)

#     pandoc_path = os.getenv('PANDOC_PATH')
#     if pandoc_path is None:
#         raise ValueError("PANDOC_PATH environment variable is not set.")

#     try:
#         # Use pandoc to convert the DOCX (with images and table placeholders) to final DOCX
#         subprocess.run([pandoc_path, "-s", temp_docx_path, "-o", output_docx_path], check=True)
#     except subprocess.CalledProcessError as e:
#         raise ValueError(f"Error during DOCX conversion: {e}")
#     finally:
#         os.remove(temp_docx_path)

def markdown_to_docx(markdown_content, output_docx_path, images):
    temp_docx_path = "temp_markdown.docx"
    
    # Convert Markdown to DOCX using pandoc
    pandoc_path = os.getenv('PANDOC_PATH')
    if pandoc_path is None:
        raise ValueError("PANDOC_PATH environment variable is not set.")
    
    with open("temp_markdown.md", "w") as f:
        f.write(markdown_content)
    
    try:
        subprocess.run([pandoc_path, "-s", "temp_markdown.md", "-o", temp_docx_path], check=True)
    except subprocess.CalledProcessError as e:
        raise ValueError(f"Error during DOCX conversion: {e}")
    finally:
        os.remove("temp_markdown.md")

    # Load the converted DOCX and add images at placeholder positions
    doc = Document(temp_docx_path)
    image_idx = 0

    for paragraph in doc.paragraphs:
        if "Image:" in paragraph.text:
            img_caption = paragraph.text.split(":")[1].strip()
            if image_idx < len(images):
                img_path = images[image_idx]
                # Replace the placeholder text with the image
                paragraph.clear()
                run = paragraph.add_run()
                run.add_picture(img_path, width=Inches(5))
                paragraph.alignment = 1  # Center alignment
                # Add caption below the image
                doc.add_paragraph(img_caption, style='Caption')
                image_idx += 1

    # Save the document with images
    doc.save(output_docx_path)
    os.remove(temp_docx_path)
    

def markdown_to_docx_with_images(markdown_content, output_docx_path, images):
    doc = Document()
    lines = markdown_content.split('\n')
    image_idx = 0

    for line in lines:
        if line.startswith("Image:"):
            img_caption = line.split(":")[1].strip()
            if image_idx < len(images):
                img_path = images[image_idx]
                # Add image to the document
                doc.add_picture(img_path, width=Inches(5))
                # Add caption below the image
                last_paragraph = doc.paragraphs[-1]
                last_paragraph.alignment = 1  # Center alignment
                doc.add_paragraph(img_caption, style='Caption')
                image_idx += 1
        elif line.startswith("Table:"):
            doc.add_paragraph(line)
        else:
            doc.add_paragraph(line)

    doc.save(output_docx_path)


def replace_placeholders_with_content(markdown_content, tables, images):
    # Replace table placeholders with actual table Markdown format
    for i, table in enumerate(tables):
        markdown_table = generate_markdown_from_table(table)
        markdown_content = markdown_content.replace(f"Table: Table_{i}", markdown_table)

    # Debugging: Output snippets of Markdown to check placeholder formatting
    print("Markdown snippet for debugging:")
    print(markdown_content[:1000])  # Print first 1000 characters of Markdown

    return markdown_content


def generate_markdown_from_table(table_data):
    if isinstance(table_data, pd.DataFrame):
        headers = list(table_data.columns)
        rows = table_data.values.tolist()
    elif isinstance(table_data, list) and len(table_data) > 0 and isinstance(table_data[0], dict):
        headers = list(table_data[0].keys())
        rows = [[row.get(header, '') for header in headers] for row in table_data]
    else:
        return "Table format error: No data or incorrect data structure."

    md_table = '| ' + ' | '.join(headers) + ' |\n'
    md_table += '| ' + ' | '.join(['---'] * len(headers)) + ' |\n'
    for row in rows:
        row_values = [str(value) for value in row]
        md_table += '| ' + ' | '.join(row_values) + ' |\n'
    
    return md_table


# Main function to execute the workflow
def main(docx_file_path, output_docx_path):
    # Create a temporary directory
    with tempfile.TemporaryDirectory() as tmpdirname:
        print(f"Using temporary directory {tmpdirname} for processing")

        # Copy the source DOCX to the temporary directory
        temp_docx_path = os.path.join(tmpdirname, 'source.docx')
        shutil.copy2(docx_file_path, temp_docx_path)

        # Process the DOCX file
        tables, images = extract_tables_and_images(temp_docx_path, tmpdirname, temp_docx_path)
        print("extracted_table", tables)

        # Convert DOCX to Markdown, adjusting paths within the temporary directory
        md_path = convert_docx_to_md_and_extract_media(temp_docx_path, tmpdirname)

        # Read and process the Markdown content
        with open(md_path, "r", encoding="utf-8") as md_file:
            markdown_content = md_file.read()
            # print("Before First prompt : ",markdown_content)

        uniform_markdown = redefine_from_markdown(markdown_content)
        # print("After First prompt : ",uniform_markdown)

        process_text = process_text_with_prompt(uniform_markdown)
        structured_content = replace_placeholders_with_content(process_text, tables, images)
        print("Markdown with structured_content:", structured_content)

        # Convert the final structured content back to a DOCX document with images embedded
        final_docx_path = os.path.join(tmpdirname, 'final_output.docx')
        # markdown_to_docx(structured_content, final_docx_path)
        markdown_to_docx(structured_content, final_docx_path, images)

        # Copy the final DOCX out of the temporary directory to the desired output location
        shutil.copy2(final_docx_path, output_docx_path)

        print("Conversion completed successfully.")


if __name__ == "__main__":
 docx_file_path=r"/Users/meghasoni/Downloads/Test7.docx"
 output_docx_path=r"/Users/meghasoni/Downloads/output_Test7.docx"
 main(docx_file_path, output_docx_path)