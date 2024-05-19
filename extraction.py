from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH  # Import WD_ALIGN_PARAGRAPH
import pandas as pd
import os


# def insert_placeholder(parent, index, placeholder_text):
#     # Create a new paragraph with the placeholder text
#     placeholder = Document().add_paragraph(placeholder_text)
#     # Insert the placeholder into the parent element at the specified index
#     parent.insert(index, placeholder._element)

# def extract_tables_and_images(docx_file, output_dir, modified_doc_path):
#     doc = Document(docx_file)
#     tables_data = []
#     images = []
#     extracted_images = set()  # To track already extracted images


#     # Ensure output directory exists
#     if not os.path.exists(output_dir):
#         os.makedirs(output_dir)

#     # Extract tables
#     for i, table in enumerate(doc.tables):
#         rows_data = []
#         prev_row_data = None  # To track the previous row for merged detection
#         for row in table.rows:
#             row_data = []
#             for cell_idx, cell in enumerate(row.cells):
#                 cell_text = cell.text.strip().replace('\n', ' ')
                
#                 # Check if this cell is part of a vertically merged set
#                 if prev_row_data and prev_row_data[cell_idx] == cell_text:
#                     row_data.append("")
#                 else:
#                     row_data.append(cell_text)
            
#             rows_data.append(row_data)
#             prev_row_data = row_data

#         # Handle possible unequal row lengths due to merged cells
#         max_len = max(len(row) for row in rows_data)
#         for row in rows_data:
#             while len(row) < max_len:
#                 row.append("")

#         df = pd.DataFrame(rows_data[1:], columns=rows_data[0])  # Keep the first row as header
#         tables_data.append(df)

#         # Find the parent of the table and its index
#         table_parent = table._element.getparent()
#         table_index = table_parent.index(table._element)

#         # Insert placeholder text for the table
#         insert_placeholder(table_parent, table_index, f"Table: Table_{i}")
        
#         # Remove the table from the document
#         table_parent.remove(table._element)

#     # Extract images
#     for rel_id, rel in doc.part.rels.items():
#         if "image" in rel.reltype and rel.target_ref not in extracted_images:
#             extracted_images.add(rel.target_ref)
#             image = rel.target_part.blob
#             image_path = os.path.join(output_dir, f'image_{len(images)}.png')
#             with open(image_path, 'wb') as img_file:
#                 img_file.write(image)
#             images.append(image_path)
    
#     print("images", images)

# # Search for image references in paragraphs and replace with placeholders
#     for paragraph in doc.paragraphs:
#         for run in paragraph.runs:
#             for rel_id in run._r.xpath('.//a:blip/@r:embed'):
#                 if rel_id in doc.part.rels and doc.part.rels[rel_id].target_ref in extracted_images:
#                     # Clear existing run and replace with placeholder text
#                     run.clear()
#                     run.text = f"Image: {doc.part.rels[rel_id].target_ref}"

#     # Save the modified document
#     doc.save(modified_doc_path)
#     return tables_data, images

import os
import pandas as pd
from docx import Document

def insert_placeholder(parent, index, placeholder_text):
    # Create a new paragraph with the placeholder text
    placeholder = Document().add_paragraph(placeholder_text)
    # Insert the placeholder into the parent element at the specified index
    parent.insert(index, placeholder._element)

def extract_tables_and_images(docx_file, output_dir, modified_doc_path):
    doc = Document(docx_file)
    tables_data = []
    images = []
    extracted_images = set()  # To track already extracted images
    image_names = {}  # To map relationship IDs to file names

    # Ensure output directory exists
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)

    # Extract tables
    for i, table in enumerate(doc.tables):
        rows_data = []
        prev_row_data = None  # To track the previous row for merged detection
        for row in table.rows:
            row_data = []
            for cell_idx, cell in enumerate(row.cells):
                cell_text = cell.text.strip().replace('\n', ' ')
                
                # Check if this cell is part of a vertically merged set
                if prev_row_data and prev_row_data[cell_idx] == cell_text:
                    row_data.append("")
                else:
                    row_data.append(cell_text)
            
            rows_data.append(row_data)
            prev_row_data = row_data

        # Handle possible unequal row lengths due to merged cells
        max_len = max(len(row) for row in rows_data)
        for row in rows_data:
            while len(row) < max_len:
                row.append("")

        df = pd.DataFrame(rows_data[1:], columns=rows_data[0])  # Keep the first row as header
        tables_data.append(df)

        # Find the parent of the table and its index
        table_parent = table._element.getparent()
        table_index = table_parent.index(table._element)

        # Insert placeholder text for the table
        insert_placeholder(table_parent, table_index, f"Table: Table_{i}")
        
        # Remove the table from the document
        table_parent.remove(table._element)

    # Extract images
    for rel_id, rel in doc.part.rels.items():
        if "image" in rel.reltype and rel.target_ref not in extracted_images:
            extracted_images.add(rel.target_ref)
            image = rel.target_part.blob
            image_name = f'image_{len(images)}.png'
            image_path = os.path.join(output_dir, image_name)
            with open(image_path, 'wb') as img_file:
                img_file.write(image)
            images.append(image_path)
            image_names[rel.target_ref] = image_name  # Map the relationship ID to the file name

    # Search for image references in paragraphs and replace with placeholders
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            for rel_id in run._r.xpath('.//a:blip/@r:embed'):
                if rel_id in doc.part.rels and doc.part.rels[rel_id].target_ref in extracted_images:
                    # Clear existing run and replace with placeholder text
                    run.clear()
                    run.text = f"Image: {image_names[doc.part.rels[rel_id].target_ref]}"

    # Save the modified document
    doc.save(modified_doc_path)
    return tables_data, images

def insert_placeholder(parent, index, text):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    p = OxmlElement('w:p')
    r = OxmlElement('w:r')
    t = OxmlElement('w:t')
    t.set(qn('xml:space'), 'preserve')
    t.text = text
    r.append(t)
    p.append(r)
    parent.insert(index, p)


# Example usage
docx_file = r"/Users/meghasoni/Downloads/Test7.docx"
output_dir = r"/Users/meghasoni/Downloads/"
modified_doc_path = r"/Users/meghasoni/Downloads/image.docx"
tables, images = extract_tables_and_images(docx_file, output_dir, modified_doc_path)
